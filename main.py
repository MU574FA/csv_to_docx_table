from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def csv_to_docx_table(csv_file_path: str, docx_file_path: str, csv_seperator: str = ",", table_start_from_right_to_left: bool = False) -> None:
    with open(csv_file_path, encoding="utf-8") as csv_file:
        # read lines of csv file...
        csv_rows = csv_file.readlines()
        # calculate the number of columens to be added to the docx table, according to number of items in a csv row.
        table_cols_num = len(csv_rows[0].strip().split(csv_seperator))
        #...
        table_rows_num = 0
        try:
            # read the already exists document
            doc   = Document(docx_file_path)
        except PackageNotFoundError:
            # initiate a new docx file...
            doc   = Document()
        # create new table...
        table = doc.add_table(table_rows_num, table_cols_num)
        # ...
        for csv_row in csv_rows:
            table_row_cells = table.add_row().cells
            # for Arabic and right-to-left languages...
            if table_start_from_right_to_left:
                table_row_cells = reversed(table_row_cells)
            for text, cell in zip(csv_row.strip().split(csv_seperator), table_row_cells):
                cell.text = text
        doc.save(docx_file_path)
